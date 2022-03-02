from datetime import datetime
import random
import os
from time import sleep
from tkinter import StringVar
from docx import Document

import pandas as pd
from hwpCtrl import HWP
# import win32com.client as win32
import PyPDF2



class Logging():

    def __init__(self, log_queue, fDir, fileType):
        self.RESULT=['OK', 'NG']
        self.dst_dir=fDir
        self.make_directory()
        log_queue.put(' Log Directory :'+ self.dst_dir+'\n')
        self.log_ftype(fileType)
        self.log_init()

    # directory가 없는 경우 생성
    def make_directory(self):
        os.makedirs(self.dst_dir, exist_ok=True)

    def log_ftype(self, fileType):
        self.fType=fileType

    def log_fDir(self, fDir):
        self.dst_dir=fDir

    def log_init(self):
        self.strResult=random.choice(self.RESULT)
        self.date_now=datetime.now().strftime('%Y%m%d %H%M%S')
        self.LogFile=f"{self.dst_dir}/{self.date_now} {self.strResult}.{self.fType}"
        # print(self.fType)
        

    def log_write(self):
        if self.fType=='txt':
            file=open(self.LogFile, 'a+', encoding="utf-8")
            file.write(self.strResult)
            file.close()
            return True
        elif self.fType=='hwp':

            hwp=HWP()
            hwp.write(self.strResult)
            hwp.savaAs(self.LogFile)
            
            return True
        elif self.fType=='xlsx':
            # self.LogFile="C:/TestDir/test2.xlsx"
            df=pd.DataFrame([self.strResult], index=['0'], columns=['0'])
            # print(df)
            df.to_excel(self.LogFile, index=False, header=False)

            return True
        elif self.fType=='docx':
            document=Document()
            document.add_paragraph(self.strResult)
            document.save(self.LogFile)
            return True
        else:
            return False
    
    def log_read(self, fName):

        # str=StringVar()
        # str.get()
        fileToRead=f"{self.dst_dir}/{fName}"
        fTypeToRead=fileToRead.split('.')[1]

        if fTypeToRead=='txt':
            file=open(fileToRead, 'r', encoding="utf-8")
            lines=file.readlines()

            file.close()
            return lines
        elif fTypeToRead=='hwp':

            hwp=HWP()
            lines=hwp.read(fileToRead)

            return lines
        

        elif fTypeToRead=='xlsx':
            df=pd.read_excel(fileToRead)

            # print(df.columns[0])
            return df.columns[0]
        elif fTypeToRead=='docx':
            document=Document(fileToRead)
            paragraphs=document.paragraphs
            lines=list()

            for para in paragraphs:
                # if "Heading" in para.style.name:
                #     print(f'{para.style.name:20}:{para.text}')
                # elif "Normal" in para.style.name:
                #     print(f'{para.style.name:20}:{para.text}')
                lines.append(para.text)

            return lines
        elif fTypeToRead=='pdf':
            file=open(fileToRead, 'rb')
            pdfReader=PyPDF2.PdfFileReader(file)

            # print(pdfReader.numPages)
            Encypted=pdfReader.isEncrypted

            if Encypted:
                pdfReader.decrypt('password')
            else:
                pass

            pdfPage=pdfReader.getPage(0)
            msg=pdfPage.extractText()

            result=list()

            print(msg)
            for txt in msg.split('\n'):
                result.append(txt)

            # rsrcmgr=PDFResourceManager()
            # retstr=StringIO()
            # laParams=LAParams()
            # device=TextConverter(rsrcmgr, retstr, laparams=laParams)
            
            # process_pdf(rsrcmgr, device, file)
            # device.close()

            # content=retstr.getvalue()
            # retstr.close()

            return result
        else:
            return False

        

    # directory내 파일 생성 및 log queue에 메시지 put
    def start_logging(self, log_queue, stop_event):   #fileType : 'txt','hwp','xlsx'
        log_queue.put('Logging started in '+self.fType+'\n')
        while True:
            try:
                self.log_init()
                if self.log_write():
                    log_queue.put(self.LogFile+' written.\n')
                else:
                    log_queue.put('log write failed.\n')
                    break
                
                sleep(1)
            # directory 소실로 인한 exception 처리
            except FileNotFoundError as err:
                log_queue.put(err)
                self.make_directory()
                log_queue.put('Creating Directory.... Start again..\n')
                break

            except Exception as e:
                print(e)
                log_queue.put(e)
                break
            
            #thead stop flag is set -> stop logging
            if stop_event.is_set():
                log_queue.put('Logging stopped.\n')
                break
            



   

    
    


    

    


    

